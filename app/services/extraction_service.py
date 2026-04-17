"""
Unified text extraction service.
- Digital PDFs  → pdfplumber (direct text layer)
- Scanned PDFs  → pypdfium2 + pytesseract (OCR fallback per page)
- Images        → pytesseract (OCR)
- DOCX          → python-docx
"""

from pathlib import Path

import pdfplumber
import pypdfium2 as pdfium
import pytesseract
from docx import Document
from PIL import Image

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

SUPPORTED_IMAGE_TYPES = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"}
SUPPORTED_PDF_TYPES = {".pdf"}
SUPPORTED_DOCX_TYPES = {".docx"}


def extract_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()

    if ext in SUPPORTED_PDF_TYPES:
        return _extract_from_pdf(file_path)
    elif ext in SUPPORTED_IMAGE_TYPES:
        return _extract_from_image(file_path)
    elif ext in SUPPORTED_DOCX_TYPES:
        return _extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_from_pdf(file_path: str) -> str:
    pages_text = []

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text and text.strip():
                pages_text.append(text.strip())
            else:
                # Scanned page — render to image and OCR
                ocr_text = _ocr_pdf_page(file_path, page.page_number - 1)
                pages_text.append(ocr_text.strip())

    return "\n\n".join(filter(None, pages_text))


def _ocr_pdf_page(file_path: str, page_index: int) -> str:
    doc = pdfium.PdfDocument(file_path)
    page = doc[page_index]
    bitmap = page.render(scale=2)          # scale=2 → 144 dpi, good OCR quality
    pil_image = bitmap.to_pil()
    return pytesseract.image_to_string(pil_image)


def _extract_from_image(file_path: str) -> str:
    return pytesseract.image_to_string(Image.open(file_path))


def _extract_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
